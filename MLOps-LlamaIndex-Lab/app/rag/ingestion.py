"""
Document ingestion pipeline.

Reads files from the uploads directory, chunks them, generates
embeddings, and upserts them into the Qdrant vector store.
"""

from pathlib import Path
from typing import List

from llama_index.core import Document, Settings as LlamaSettings, VectorStoreIndex, StorageContext
from llama_index.core.node_parser import SentenceSplitter

from app.core.config import settings
from app.core.logging import get_logger
from app.rag.embeddings import build_embedding_model
from app.rag.vector_store import get_vector_store

logger = get_logger(__name__)

# Supported file extensions and their simple readers
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def _read_text_file(path: Path) -> str:
    """Read a plain-text or markdown file."""
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf_file(path: Path) -> str:
    """Extract text from a PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx_file(path: Path) -> str:
    """Extract text from a .docx file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def load_document(path: Path) -> Document:
    """Load a single file into a LlamaIndex Document."""
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        text = _read_text_file(path)
    elif ext == ".pdf":
        text = _read_pdf_file(path)
    elif ext == ".docx":
        text = _read_docx_file(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return Document(
        text=text,
        metadata={"source": path.name, "file_path": str(path)},
    )


def load_all_documents(directory: Path | None = None) -> List[Document]:
    """Load every supported file from *directory* (defaults to upload_path)."""
    directory = directory or settings.upload_path
    docs: List[Document] = []
    for file in sorted(directory.iterdir()):
        if file.suffix.lower() in SUPPORTED_EXTENSIONS:
            logger.info("Loading document: %s", file.name)
            docs.append(load_document(file))
    logger.info("Loaded %d document(s) from %s", len(docs), directory)
    return docs


def ingest_documents(documents: List[Document] | None = None) -> VectorStoreIndex:
    """
    Run the full ingestion pipeline:
    1. Load documents (if not provided).
    2. Chunk with SentenceSplitter.
    3. Embed with the configured model.
    4. Store in Qdrant.

    Returns the resulting VectorStoreIndex.
    """
    if documents is None:
        documents = load_all_documents()

    if not documents:
        logger.warning("No documents to ingest.")
        # Return an empty index so callers don't crash
        embed_model = build_embedding_model()
        LlamaSettings.embed_model = embed_model
        vector_store = get_vector_store()
        storage_ctx = StorageContext.from_defaults(vector_store=vector_store)
        return VectorStoreIndex([], storage_context=storage_ctx)

    # Configure chunking
    splitter = SentenceSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )

    # Configure embedding model globally for LlamaIndex
    embed_model = build_embedding_model()
    LlamaSettings.embed_model = embed_model

    # Build the vector store index (automatically chunks → embeds → stores)
    vector_store = get_vector_store()
    storage_ctx = StorageContext.from_defaults(vector_store=vector_store)

    logger.info(
        "Ingesting %d document(s)  chunk_size=%d  overlap=%d",
        len(documents), settings.chunk_size, settings.chunk_overlap,
    )
    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_ctx,
        transformations=[splitter],
        show_progress=True,
    )
    logger.info("Ingestion complete.")
    return index
